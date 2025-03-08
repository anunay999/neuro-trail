import logging
import io
from typing import Dict, Any, List, Optional
import docx

from app.core.plugin_base import DocumentPlugin
from app.utils.text_utils import split_text, extract_metadata

# Configure logging
logger = logging.getLogger(__name__)


class DOCXHandler(DocumentPlugin):
    """Plugin for handling DOCX documents"""
    
    plugin_type = "document_handler"
    plugin_name = "docx_handler"
    plugin_version = "0.1.0"
    plugin_description = "Handler for DOCX documents"
    
    def __init__(self):
        """Initialize DOCX handler"""
        self.initialized = False
    
    async def initialize(self, **kwargs) -> bool:
        """
        Initialize the plugin
        
        Args:
            **kwargs: Additional parameters
            
        Returns:
            bool: True if initialization successful
        """
        self.initialized = True
        logger.info("DOCX handler initialized")
        return True
    
    async def shutdown(self) -> bool:
        """
        Shutdown the plugin
        
        Returns:
            bool: True if shutdown successful
        """
        self.initialized = False
        logger.info("DOCX handler shutdown")
        return True
    
    async def process_document(
        self, 
        file_content: bytes, 
        file_name: str, 
        **kwargs
    ) -> Dict[str, Any]:
        """
        Process a DOCX document
        
        Args:
            file_content: Binary content of the DOCX file
            file_name: Name of the file
            **kwargs: Additional parameters
            
        Returns:
            Dict[str, Any]: Dictionary with processed document data
        """
        logger.info(f"Processing DOCX document: {file_name}")
        
        try:
            # Create a BytesIO object from file_content
            file_stream = io.BytesIO(file_content)
            
            # Load the document
            doc = docx.Document(file_stream)
            
            # Extract title from document or use filename
            title = self._extract_title(doc) or file_name
            
            # Extract metadata
            metadata = {
                "title": title,
                "author": self._extract_author(doc) or "Unknown Author",
                "format": "docx"
            }
            
            # Extract chapters
            chapters = self._extract_chapters(doc)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            # Combine all paragraphs
            full_text = "\n\n".join(paragraphs)
            
            # Split text into chunks
            text_chunks = split_text(full_text, chunk_size=1000, overlap=200)
            
            logger.info(f"Processed DOCX document: {title}, {len(text_chunks)} chunks")
            
            # Return processed document data
            return {
                "metadata": metadata,
                "chapters": chapters,
                "texts": text_chunks,
                "full_text": full_text
            }
            
        except Exception as e:
            logger.exception(f"Error processing DOCX document: {e}")
            raise Exception(f"Error processing DOCX document: {str(e)}")
    
    def _extract_title(self, doc) -> Optional[str]:
        """
        Extract title from document
        
        Args:
            doc: python-docx Document
            
        Returns:
            Optional[str]: Title if found, None otherwise
        """
        # Try to get title from document properties
        if doc.core_properties.title:
            return doc.core_properties.title
        
        # Try to get title from first heading
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading 1') and para.text.strip():
                return para.text.strip()
            
        return None
    
    def _extract_author(self, doc) -> Optional[str]:
        """
        Extract author from document
        
        Args:
            doc: python-docx Document
            
        Returns:
            Optional[str]: Author if found, None otherwise
        """
        # Try to get author from document properties
        if doc.core_properties.author:
            return doc.core_properties.author
            
        return None
    
    def _extract_chapters(self, doc) -> List[Dict[str, Any]]:
        """
        Extract chapters from document headings
        
        Args:
            doc: python-docx Document
            
        Returns:
            List[Dict[str, Any]]: List of chapters
        """
        chapters = []
        
        # Find all heading paragraphs
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading 1') and para.text.strip():
                chapters.append({
                    "title": para.text.strip(),
                    "seq": len(chapters) + 1
                })
        
        # If no headings found, try to infer structure
        if not chapters:
            chapter_count = 0
            current_title = None
            
            for para in doc.paragraphs:
                # Look for potential chapter markers
                text = para.text.strip()
                if text and (text.startswith("Chapter ") or 
                            text.startswith("CHAPTER ") or
                            (len(text) < 50 and text.isupper())):
                    current_title = text
                    chapter_count += 1
                    chapters.append({
                        "title": current_title,
                        "seq": chapter_count
                    })
        
        return chapters